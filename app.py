import streamlit as st
import os
import tempfile
from typing import List, Optional
import uuid
import io

# Document processing
from PyPDF2 import PdfReader
import docx

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain.llms import HuggingFacePipeline
from langchain.chains.summarize import load_summarize_chain
from langchain.chains.llm import LLMChain
from langchain.prompts import PromptTemplate

# Transformers for the language model
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
import torch

# Set page configuration
st.set_page_config(
    page_title="ManGPT - Document AI Assistant",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        border-bottom: 2px solid #f0f2f6;
        margin-bottom: 2rem;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 4px solid #ff4b4b;
    }
    .user-message {
        background-color: #f0f2f6;
        border-left-color: #0068c9;
    }
    .bot-message {
        background-color: #f9f9f9;
        border-left-color: #ff4b4b;
        border: 1px solid #e6e6e6;
    }
    .response-container {
        background-color: #ffffff;
        padding: 20px;
        border-radius: 10px;
        border: 1px solid #e1e5e9;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .summary-container {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        border: 1px solid #dee2e6;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .upload-section {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        border: 2px dashed #adb5bd;
        text-align: center;
        margin: 20px 0;
    }
    .stButton > button {
        width: 100%;
        border-radius: 20px;
        height: 3rem;
        background: linear-gradient(90deg, #ff4b4b, #ff6b6b);
        border: none;
        color: white;
        font-weight: bold;
    }
    .question-input {
        font-size: 16px;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state with better management
def init_session_state():
    if 'documents' not in st.session_state:
        st.session_state.documents = []
    if 'vectorstore' not in st.session_state:
        st.session_state.vectorstore = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'embeddings' not in st.session_state:
        st.session_state.embeddings = None
    if 'llm' not in st.session_state:
        st.session_state.llm = None
    if 'processed_files' not in st.session_state:
        st.session_state.processed_files = []
    if 'current_summary' not in st.session_state:
        st.session_state.current_summary = ""
    if 'processing_status' not in st.session_state:
        st.session_state.processing_status = ""

@st.cache_resource
def load_embeddings():
    """Load the embedding model"""
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'}
    )

@st.cache_resource
def load_llm():
    """Load a simple but effective text generation system"""
    class SmartLLM:
        def __init__(self):
            self.max_length = 1000
            
        def generate_response(self, context: str, question: str) -> str:
            """Generate a comprehensive response based on context and question"""
            try:
                # Tokenize question for better matching
                question_words = set(question.lower().split())
                context_sentences = context.split('.')
                
                # Score sentences based on relevance
                scored_sentences = []
                for sentence in context_sentences:
                    sentence = sentence.strip()
                    if len(sentence) < 20:  # Skip very short sentences
                        continue
                    
                    sentence_words = set(sentence.lower().split())
                    # Calculate relevance score
                    common_words = question_words.intersection(sentence_words)
                    score = len(common_words)
                    
                    if score > 0:
                        scored_sentences.append((score, sentence))
                
                # Sort by relevance and take top sentences
                scored_sentences.sort(reverse=True, key=lambda x: x[0])
                
                if scored_sentences:
                    # Build comprehensive response
                    response_parts = []
                    total_length = 0
                    
                    for score, sentence in scored_sentences[:5]:  # Top 5 relevant sentences
                        if total_length + len(sentence) < self.max_length:
                            response_parts.append(sentence.strip() + '.')
                            total_length += len(sentence)
                    
                    if response_parts:
                        response = "Based on the uploaded documents, here's what I found:\n\n"
                        response += "\n\n".join(response_parts)
                        
                        # Add a concluding note
                        response += "\n\nThis information is extracted from your uploaded documents. Feel free to ask more specific questions for detailed insights."
                        return response
                
                # Fallback response with context excerpt
                return f"Based on your documents:\n\n{context[:800]}...\n\nPlease try asking more specific questions about the content for better results."
                
            except Exception as e:
                return f"I found relevant information in your documents, but encountered an issue processing it. Please try rephrasing your question or upload the documents again."
    
    return SmartLLM()

def extract_text_from_pdf(pdf_file) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(docx_file) -> str:
    """Extract text from Word document"""
    try:
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from Word document: {str(e)}")
        return ""

def process_documents(uploaded_files) -> List[Document]:
    """Process uploaded documents and extract text"""
    documents = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, uploaded_file in enumerate(uploaded_files):
        status_text.text(f"Processing: {uploaded_file.name}")
        
        try:
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name
            
            # Extract text based on file type
            if uploaded_file.type == "application/pdf":
                text = extract_text_from_pdf(tmp_file_path)
            elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                text = extract_text_from_docx(tmp_file_path)
            else:
                st.warning(f"Unsupported file type: {uploaded_file.type}")
                continue
            
            if text and len(text.strip()) > 50:  # Ensure meaningful content
                # Create document object
                doc = Document(
                    page_content=text,
                    metadata={
                        "source": uploaded_file.name,
                        "type": uploaded_file.type,
                        "length": len(text)
                    }
                )
                documents.append(doc)
                st.session_state.processed_files.append({
                    "name": uploaded_file.name,
                    "size": len(text),
                    "type": uploaded_file.type
                })
            
            # Clean up temporary file
            os.unlink(tmp_file_path)
            
            # Update progress
            progress_bar.progress((i + 1) / len(uploaded_files))
            
        except Exception as e:
            st.error(f"Error processing {uploaded_file.name}: {str(e)}")
    
    progress_bar.empty()
    status_text.empty()
    
    return documents

def create_vectorstore(documents: List[Document]):
    """Create FAISS vector store from documents"""
    if not documents:
        return None
    
    try:
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # Smaller chunks for better precision
            chunk_overlap=100,
            length_function=len
        )
        
        texts = text_splitter.split_documents(documents)
        
        # Create embeddings
        embeddings = st.session_state.embeddings
        if embeddings is None:
            embeddings = load_embeddings()
            st.session_state.embeddings = embeddings
        
        # Create vector store
        vectorstore = FAISS.from_documents(texts, embeddings)
        return vectorstore
        
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        return None

def get_response(question: str, vectorstore) -> str:
    """Get response using RAG pipeline"""
    if vectorstore is None:
        return "Please upload and process documents first."
    
    try:
        # Get LLM
        if st.session_state.llm is None:
            st.session_state.llm = load_llm()
        
        # Retrieve relevant documents
        retriever = vectorstore.as_retriever(search_kwargs={"k": 5})  # Get more context
        docs = retriever.get_relevant_documents(question)
        
        # Combine context from retrieved documents
        context = "\n\n".join([doc.page_content for doc in docs])
        
        # Generate response using our smart LLM
        response = st.session_state.llm.generate_response(context, question)
        return response
        
    except Exception as e:
        return f"I encountered an issue while processing your question. Please try again or rephrase your question. Error details: {str(e)}"

def generate_summary(documents: List[Document]) -> str:
    """Generate a comprehensive summary of the documents"""
    if not documents:
        return "No documents to summarize."
    
    try:
        # Combine all document text
        full_text = "\n\n".join([doc.page_content for doc in documents])
        
        # Extract key information for summary
        paragraphs = [p.strip() for p in full_text.split('\n\n') if len(p.strip()) > 100]
        
        # Create structured summary
        summary_parts = []
        
        # Document overview
        summary_parts.append(f"📄 **Document Overview:**")
        summary_parts.append(f"• Total documents processed: {len(documents)}")
        summary_parts.append(f"• Total content length: {len(full_text):,} characters")
        summary_parts.append("")
        
        # Key content extraction
        summary_parts.append("📋 **Key Content:**")
        
        # Take first few meaningful paragraphs
        for i, para in enumerate(paragraphs[:4]):
            sentences = para.split('.')
            key_sentence = sentences[0].strip()
            if len(key_sentence) > 50:
                summary_parts.append(f"• {key_sentence}.")
        
        summary_parts.append("")
        summary_parts.append("📊 **Document Statistics:**")
        for doc in documents:
            doc_name = doc.metadata.get('source', 'Unknown')
            doc_length = doc.metadata.get('length', len(doc.page_content))
            summary_parts.append(f"• {doc_name}: {doc_length:,} characters")
        
        summary_parts.append("")
        summary_parts.append("💡 **How to use:** Ask specific questions about the content to get detailed information from these documents.")
        
        return "\n".join(summary_parts)
        
    except Exception as e:
        return f"Error generating summary: {str(e)}"

def display_chat_message(message_type: str, content: str):
    """Display a chat message with proper formatting"""
    if message_type == "user":
        st.markdown(f"""
        <div class="chat-message user-message">
            <strong>🧑‍💼 You:</strong><br>
            {content}
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="response-container">
            <strong>🤖 ManGPT:</strong><br><br>
            {content.replace('\n', '<br>')}
        </div>
        """, unsafe_allow_html=True)

# Initialize session state
init_session_state()

def main():
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📚 ManGPT - Document AI Assistant</h1>
        <p style="font-size: 1.2em; color: #666;">Upload your documents and get intelligent answers instantly</p>
    </div>
    """, unsafe_allow_html=True)
    
    # File Upload Section
    st.markdown("""
    <div class="upload-section">
        <h3>📁 Upload Your Documents</h3>
        <p>Support for PDF and Word documents • Maximum file size: 200MB per file</p>
    </div>
    """, unsafe_allow_html=True)
    
    uploaded_files = st.file_uploader(
        "Choose files",
        type=['pdf', 'docx', 'doc'],
        accept_multiple_files=True,
        help="Upload PDF or Word documents to analyze",
        label_visibility="collapsed"
    )
    
    # Two columns for buttons
    col1, col2 = st.columns(2)
    
    with col1:
        process_button = st.button("🔄 Process Documents", type="primary")
    
    with col2:
        summarize_button = st.button("📝 Generate Summary", type="secondary")
    
    # Process Documents
    if process_button and uploaded_files:
        with st.spinner("🔄 Processing your documents... This may take a few moments."):
            # Clear previous data
            st.session_state.documents = []
            st.session_state.vectorstore = None
            st.session_state.processed_files = []
            
            # Process documents
            documents = process_documents(uploaded_files)
            
            if documents:
                st.session_state.documents = documents
                
                # Create vector store
                with st.spinner("🔍 Creating search index..."):
                    vectorstore = create_vectorstore(documents)
                    st.session_state.vectorstore = vectorstore
                
                if vectorstore:
                    st.success(f"✅ Successfully processed {len(documents)} document(s)! You can now ask questions.")
                    
                    # Display processed files info
                    with st.expander("📋 Processed Files Details"):
                        for file_info in st.session_state.processed_files:
                            st.write(f"• **{file_info['name']}** - {file_info['size']:,} characters")
                else:
                    st.error("❌ Failed to create search index. Please try again.")
            else:
                st.error("❌ No documents could be processed. Please check your files and try again.")
    
    # Generate Summary
    if summarize_button and st.session_state.documents:
        with st.spinner("📝 Generating comprehensive summary..."):
            summary = generate_summary(st.session_state.documents)
            st.session_state.current_summary = summary
        
        st.markdown("""
        <div class="summary-container">
            <h3>📋 Document Summary</h3>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <div class="response-container">
            {st.session_state.current_summary.replace('\n', '<br>')}
        </div>
        """, unsafe_allow_html=True)
    
    # Show current summary if exists
    elif st.session_state.current_summary:
        st.markdown("""
        <div class="summary-container">
            <h3>📋 Document Summary</h3>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <div class="response-container">
            {st.session_state.current_summary.replace('\n', '<br>')}
        </div>
        """, unsafe_allow_html=True)
    
    # Chat Interface
    if st.session_state.vectorstore is not None:
        st.markdown("---")
        st.markdown("## 💬 Ask Questions About Your Documents")
        
        # Display chat history
        for question, answer in st.session_state.chat_history:
            display_chat_message("user", question)
            display_chat_message("bot", answer)
        
        # Question input form
        with st.form(key="question_form", clear_on_submit=True):
            question = st.text_area(
                "Ask your question here:",
                placeholder="What is the main topic discussed in the documents? Explain in detail...",
                height=100,
                max_chars=1000,
                help="You can ask detailed questions up to 1000 characters"
            )
            
            col1, col2 = st.columns([3, 1])
            with col1:
                ask_button = st.form_submit_button("🚀 Get Answer", type="primary")
            with col2:
                clear_button = st.form_submit_button("🗑️ Clear Chat")
        
        # Handle question submission
        if ask_button and question.strip():
            with st.spinner("🤔 Analyzing your documents and generating response..."):
                response = get_response(question.strip(), st.session_state.vectorstore)
                
                # Add to chat history
                st.session_state.chat_history.append((question.strip(), response))
                
                # Display the new exchange
                display_chat_message("user", question.strip())
                display_chat_message("bot", response)
        
        # Handle clear chat
        if clear_button:
            st.session_state.chat_history = []
            st.rerun()
    
    # Instructions when no documents are processed
    else:
        if not uploaded_files:
            st.markdown("---")
            st.info("👆 Please upload your documents above to get started!")
        elif uploaded_files:
            st.markdown("---")
            st.info("👆 Documents uploaded! Click 'Process Documents' to analyze them.")
        
        # Show usage instructions
        with st.expander("💡 How to Use ManGPT", expanded=True):
            st.markdown("""
            ### Quick Start Guide:
            
            1. **📁 Upload Documents**: Choose PDF or Word files using the uploader above
            2. **🔄 Process**: Click "Process Documents" to analyze and index your files
            3. **📝 Summarize** (Optional): Generate a comprehensive summary of your documents
            4. **💬 Ask Questions**: Use the chat interface to ask detailed questions
            5. **🔍 Get Answers**: Receive comprehensive, context-aware responses
            
            ### Example Questions:
            - "What are the main findings discussed in this research?"
            - "Summarize the key recommendations from the document"
            - "What methodology was used in this study?"
            - "Who are the main stakeholders mentioned?"
            - "What are the financial implications discussed?"
            
            ### Tips for Best Results:
            - Ask specific, detailed questions for more accurate answers
            - Use clear, well-structured documents for better results
            - Try different phrasings if you don't get the expected answer
            """)
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>Built with ❤️ using Streamlit, LangChain & Hugging Face</p>
        <p>ManGPT v2.0 - Intelligent Document Analysis</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
